"""
Generate Use Case Description tables for Hotel Management System.
Each UC is an individual table (not grouped as 'Manage X').
Output: Word document with all UC tables.
"""
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── UC Data ──────────────────────────────────────────────────────────

UC_DATA = [
    # ═══════════════════ AUTHENTICATION (Guest/Customer) ═══════════════════
    {
        "id": "UC-01", "name": "Register Account",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Guest", "secondary_actor": "System",
        "description": "Allows a guest to create a new customer account by providing personal information and credentials.",
        "trigger": "Guest decides to create an account to book rooms.",
        "preconditions": "PRE-1: Guest is not logged in.\nPRE-2: Guest has a valid email address.",
        "postconditions": "POST-1: A new customer account is created and stored.\nPOST-2: Guest can log in with the new credentials.",
        "normal_flow": (
            "1. Guest selects the registration option.\n"
            "2. System requests personal information (full name, email, phone, address) and password.\n"
            "3. Guest provides required information.\n"
            "4. System validates the information and checks email uniqueness.\n"
            "5. System creates the account and assigns Customer role.\n"
            "6. System confirms successful registration."
        ),
        "alternative_flows": "AF-1 (at step 3): Guest chooses to register via Google OAuth. System redirects to Google for authentication, retrieves profile, creates account.",
        "exceptions": "EF-1 (at step 4): Email already exists. System notifies guest and requests a different email.\nEF-2 (at step 4): Invalid data format. System highlights errors and requests correction.",
        "priority": "High", "frequency": "Medium",
        "business_rules": "BR-1: Email must be unique across all accounts.\nBR-2: Password must meet minimum security requirements.",
        "other_info": "", "assumptions": "ASM-1: Email service is available for verification if needed."
    },
    {
        "id": "UC-02", "name": "Login",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Guest", "secondary_actor": "System",
        "description": "Allows a user to authenticate using email/password or Google OAuth to access the system.",
        "trigger": "User wants to access authenticated features.",
        "preconditions": "PRE-1: User has a registered account.\nPRE-2: User is not currently logged in.",
        "postconditions": "POST-1: User session is established.\nPOST-2: User is redirected to the appropriate dashboard based on role.",
        "normal_flow": (
            "1. User provides email and password.\n"
            "2. System validates credentials.\n"
            "3. System creates a session and determines user role.\n"
            "4. System redirects user to the corresponding area (Admin/Staff/Customer)."
        ),
        "alternative_flows": "AF-1 (at step 1): User selects Google OAuth login. System redirects to Google, receives token, matches or creates account.",
        "exceptions": "EF-1 (at step 2): Invalid credentials. System notifies user.\nEF-2 (at step 2): Account is deactivated. System notifies user.",
        "priority": "High", "frequency": "Very High",
        "business_rules": "BR-1: Password is verified using BCrypt hashing.\nBR-2: Account must be active to login.",
        "other_info": "", "assumptions": "ASM-1: Google OAuth service is available."
    },
    {
        "id": "UC-03", "name": "Forgot Password",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Allows a customer to request a password reset via OTP sent to their registered email.",
        "trigger": "Customer cannot remember their password.",
        "preconditions": "PRE-1: Customer has a registered account with valid email.",
        "postconditions": "POST-1: An OTP is sent to the customer's email.\nPOST-2: Customer can proceed to verify OTP and reset password.",
        "normal_flow": (
            "1. Customer selects forgot password option.\n"
            "2. System requests the registered email address.\n"
            "3. Customer provides email.\n"
            "4. System verifies email exists and generates an OTP.\n"
            "5. System sends OTP to the email.\n"
            "6. System confirms OTP has been sent."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 4): Email not found. System notifies customer.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: OTP expires after a configured time period.",
        "other_info": "", "assumptions": "ASM-1: Email service is operational."
    },
    {
        "id": "UC-04", "name": "Verify OTP",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer verifies identity by entering the OTP received via email.",
        "trigger": "Customer has received an OTP for password reset.",
        "preconditions": "PRE-1: OTP has been sent to customer's email.\nPRE-2: OTP has not expired.",
        "postconditions": "POST-1: Customer identity is verified.\nPOST-2: Customer is allowed to set a new password.",
        "normal_flow": (
            "1. System requests the OTP code.\n"
            "2. Customer enters the OTP.\n"
            "3. System validates the OTP.\n"
            "4. System confirms verification and allows password reset."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Invalid or expired OTP. System notifies customer.\nEF-2 (at step 3): Maximum attempts exceeded. System blocks further attempts.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: OTP is single-use.\nBR-2: OTP has a time-to-live limit.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-05", "name": "Reset Password",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer sets a new password after successful OTP verification.",
        "trigger": "Customer has successfully verified OTP.",
        "preconditions": "PRE-1: Customer has passed OTP verification.",
        "postconditions": "POST-1: Customer's password is updated.\nPOST-2: Customer can login with the new password.",
        "normal_flow": (
            "1. System requests new password and confirmation.\n"
            "2. Customer provides new password and confirms it.\n"
            "3. System validates password strength and match.\n"
            "4. System updates the password (hashed).\n"
            "5. System confirms password reset."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Passwords do not match. System requests re-entry.\nEF-2 (at step 3): Password does not meet requirements. System notifies.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Password is stored using BCrypt hashing.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-06", "name": "Change Password",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Authenticated customer changes their current password.",
        "trigger": "Customer wants to update their password for security.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: Password is updated.",
        "normal_flow": (
            "1. Customer provides current password, new password, and confirmation.\n"
            "2. System verifies current password.\n"
            "3. System validates new password strength and match.\n"
            "4. System updates password.\n"
            "5. System confirms change."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Current password incorrect. System notifies.\nEF-2 (at step 3): New password same as old. System rejects.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: New password must differ from current password.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-07", "name": "Complete Profile",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer who registered via Google OAuth completes missing profile information.",
        "trigger": "Customer logged in via Google OAuth with incomplete profile.",
        "preconditions": "PRE-1: Customer registered via Google OAuth.\nPRE-2: Profile has missing required fields (phone, address).",
        "postconditions": "POST-1: Customer profile is complete.\nPOST-2: Customer can use all system features.",
        "normal_flow": (
            "1. System detects incomplete profile and requests missing information.\n"
            "2. Customer provides phone number and address.\n"
            "3. System validates and saves the information.\n"
            "4. System confirms profile completion."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Invalid phone format. System requests correction.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Phone number must be valid format.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ROOM BROWSING (Customer/Guest) ═══════════════════
    {
        "id": "UC-08", "name": "Browse Available Rooms",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Guest/Customer", "secondary_actor": "System",
        "description": "User browses the list of available rooms with filtering options.",
        "trigger": "User wants to view room options before booking.",
        "preconditions": "PRE-1: System has rooms configured.",
        "postconditions": "POST-1: User sees a list of available rooms with details and pricing.",
        "normal_flow": (
            "1. User accesses the room listing.\n"
            "2. System retrieves and displays available rooms with type, price, capacity, and images.\n"
            "3. User optionally filters by room type, capacity, or date range.\n"
            "4. System updates the listing based on filters."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): No rooms available. System displays appropriate message.",
        "priority": "High", "frequency": "Very High",
        "business_rules": "BR-1: Only rooms with available status are shown to guests.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-09", "name": "View Room Detail",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Guest/Customer", "secondary_actor": "System",
        "description": "User views detailed information about a specific room type including amenities and images.",
        "trigger": "User selects a room from the listing.",
        "preconditions": "PRE-1: Room exists in the system.",
        "postconditions": "POST-1: User sees full room details (type, price, capacity, amenities, images, description).",
        "normal_flow": (
            "1. User selects a specific room.\n"
            "2. System retrieves room details including type info, amenities, and images.\n"
            "3. System displays the detailed information."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Room not found. System notifies user.",
        "priority": "High", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ BOOKING (Customer) ═══════════════════
    {
        "id": "UC-10", "name": "Create Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer creates a new room booking by selecting dates and room type.",
        "trigger": "Customer decides to book a room.",
        "preconditions": "PRE-1: Customer is logged in.\nPRE-2: Rooms of selected type are available for the date range.",
        "postconditions": "POST-1: Booking record is created with Pending status.\nPOST-2: Total price and deposit amount are calculated.",
        "normal_flow": (
            "1. Customer selects room type, check-in date, and check-out date.\n"
            "2. System checks room availability for the date range.\n"
            "3. System calculates total price and deposit amount.\n"
            "4. Customer optionally applies a voucher code.\n"
            "5. Customer selects payment type (online deposit or pay at hotel).\n"
            "6. System creates the booking and assigns a room.\n"
            "7. System confirms the booking with details."
        ),
        "alternative_flows": "AF-1 (at step 4): Customer applies voucher. System validates voucher and recalculates price with discount.",
        "exceptions": "EF-1 (at step 2): No rooms available for selected dates. System suggests alternative dates.\nEF-2 (at step 4): Invalid or expired voucher. System notifies customer.",
        "priority": "High", "frequency": "High",
        "business_rules": "BR-1: Deposit is calculated as a percentage of total price based on room type.\nBR-2: Booking date cannot be in the past.",
        "other_info": "", "assumptions": "ASM-1: Room pricing is pre-configured per room type."
    },
    {
        "id": "UC-11", "name": "Confirm Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer reviews and confirms booking details before proceeding to payment.",
        "trigger": "Booking details have been calculated.",
        "preconditions": "PRE-1: Booking calculation is complete.",
        "postconditions": "POST-1: Booking is confirmed and awaiting payment.\nPOST-2: Customer is directed to payment if online deposit selected.",
        "normal_flow": (
            "1. System displays booking summary (room, dates, price, deposit, voucher discount).\n"
            "2. Customer reviews the details.\n"
            "3. Customer confirms the booking.\n"
            "4. System saves the booking record.\n"
            "5. If online payment selected, system redirects to payment."
        ),
        "alternative_flows": "AF-1 (at step 3): Customer cancels. System discards the booking draft.",
        "exceptions": "",
        "priority": "High", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-12", "name": "View Booking Status",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views the current status of a specific booking.",
        "trigger": "Customer wants to check booking progress.",
        "preconditions": "PRE-1: Customer is logged in.\nPRE-2: Booking exists.",
        "postconditions": "POST-1: Customer sees current booking status and details.",
        "normal_flow": (
            "1. Customer accesses their booking.\n"
            "2. System retrieves booking details and current status.\n"
            "3. System displays status (Pending/Confirmed/Checked-in/Checked-out/Cancelled)."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-13", "name": "Extend Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer requests to extend their current stay by additional nights.",
        "trigger": "Customer wants to stay longer than originally booked.",
        "preconditions": "PRE-1: Customer has an active (checked-in) booking.\nPRE-2: Room is available for the extension period.",
        "postconditions": "POST-1: Booking extension is recorded.\nPOST-2: Additional charges are calculated.",
        "normal_flow": (
            "1. Customer selects the booking to extend.\n"
            "2. Customer specifies new check-out date.\n"
            "3. System checks room availability for extended period.\n"
            "4. System calculates additional cost.\n"
            "5. Customer confirms the extension.\n"
            "6. System records the extension and updates booking."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Room not available for extension period. System notifies customer.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Extension price is based on room type hourly/daily rate.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-14", "name": "View Booking History",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views list of all past and current bookings.",
        "trigger": "Customer wants to review their booking history.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: Customer sees a list of all their bookings.",
        "normal_flow": (
            "1. Customer accesses booking history.\n"
            "2. System retrieves all bookings for the customer.\n"
            "3. System displays bookings with dates, room, status, and price."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): No bookings found. System displays empty state message.",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-15", "name": "View Booking Detail",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views detailed information of a specific booking.",
        "trigger": "Customer selects a booking from their history.",
        "preconditions": "PRE-1: Customer is logged in.\nPRE-2: Booking belongs to the customer.",
        "postconditions": "POST-1: Full booking details are displayed.",
        "normal_flow": (
            "1. Customer selects a booking.\n"
            "2. System retrieves full details (room, dates, price, payment status, extensions, occupants).\n"
            "3. System displays all booking information."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "BR-1: Customer can only view their own bookings.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ PAYMENT (Customer) ═══════════════════
    {
        "id": "UC-16", "name": "Make Online Payment",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "VNPay Gateway",
        "description": "Customer makes an online deposit payment via VNPay for their booking.",
        "trigger": "Customer confirms booking with online payment option.",
        "preconditions": "PRE-1: Booking is confirmed and pending payment.\nPRE-2: VNPay gateway is operational.",
        "postconditions": "POST-1: Payment is processed and recorded.\nPOST-2: Booking status is updated to Confirmed.",
        "normal_flow": (
            "1. System generates VNPay payment URL with booking amount.\n"
            "2. Customer is redirected to VNPay gateway.\n"
            "3. Customer completes payment on VNPay.\n"
            "4. VNPay sends payment result (IPN callback).\n"
            "5. System verifies payment signature and records transaction.\n"
            "6. System updates booking status to Confirmed.\n"
            "7. System displays payment success."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Customer cancels payment. System shows failed status.\nEF-2 (at step 5): Payment verification fails. System rejects and logs.",
        "priority": "High", "frequency": "High",
        "business_rules": "BR-1: Payment amount must match booking deposit.\nBR-2: VNPay signature must be verified.",
        "other_info": "", "assumptions": "ASM-1: VNPay gateway is available and configured."
    },

    # ═══════════════════ CUSTOMER PROFILE & FEEDBACK ═══════════════════
    {
        "id": "UC-17", "name": "View Profile",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views their personal profile information.",
        "trigger": "Customer wants to review personal information.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: Profile information is displayed.",
        "normal_flow": (
            "1. Customer accesses profile section.\n"
            "2. System retrieves and displays account details (name, email, phone, address, membership level, loyalty points)."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-18", "name": "Edit Profile",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer updates their personal profile information.",
        "trigger": "Customer wants to update personal details.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: Profile information is updated.",
        "normal_flow": (
            "1. Customer accesses profile editing.\n"
            "2. System displays current information in editable form.\n"
            "3. Customer modifies desired fields (name, phone, address).\n"
            "4. System validates updated information.\n"
            "5. System saves changes and confirms."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 4): Invalid data. System highlights errors.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Email cannot be changed (used as login identifier).",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-19", "name": "Submit Feedback",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer submits a review/feedback for a completed booking.",
        "trigger": "Customer wants to share their experience after a stay.",
        "preconditions": "PRE-1: Customer is logged in.\nPRE-2: Customer has a completed (checked-out) booking.",
        "postconditions": "POST-1: Feedback is recorded in the system.",
        "normal_flow": (
            "1. Customer selects a completed booking to review.\n"
            "2. System requests rating and comments.\n"
            "3. Customer provides rating and feedback text.\n"
            "4. System validates and saves the feedback.\n"
            "5. System confirms submission."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Booking already reviewed. System notifies customer.",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "BR-1: One feedback per booking.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-20", "name": "View My Reviews",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views all their submitted feedback/reviews.",
        "trigger": "Customer wants to see their past reviews.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: List of customer's reviews is displayed.",
        "normal_flow": (
            "1. Customer accesses reviews section.\n"
            "2. System retrieves all feedback submitted by the customer.\n"
            "3. System displays reviews with ratings, comments, and booking references."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ SERVICE REQUESTS (Customer) ═══════════════════
    {
        "id": "UC-21", "name": "Submit Service Request",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer submits a service request (housekeeping, maintenance, etc.) during an active stay.",
        "trigger": "Customer needs hotel services during stay.",
        "preconditions": "PRE-1: Customer is logged in.\nPRE-2: Customer has an active (checked-in) booking.",
        "postconditions": "POST-1: Service request is created with Pending status.\nPOST-2: Staff is notified of the request.",
        "normal_flow": (
            "1. Customer selects the active booking.\n"
            "2. Customer specifies service type, description, and priority.\n"
            "3. System validates the request.\n"
            "4. System creates the service request with Pending status.\n"
            "5. System confirms submission."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): No active booking found. System notifies customer.",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "BR-1: Service types include: housekeeping, maintenance, room service, wake-up call, etc.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-22", "name": "View My Service Requests",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Customer", "secondary_actor": "System",
        "description": "Customer views all their service requests and their current status.",
        "trigger": "Customer wants to track service request progress.",
        "preconditions": "PRE-1: Customer is logged in.",
        "postconditions": "POST-1: List of service requests with statuses is displayed.",
        "normal_flow": (
            "1. Customer accesses service requests section.\n"
            "2. System retrieves all requests for the customer.\n"
            "3. System displays requests with type, status, and timestamps."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - BOOKING MANAGEMENT ═══════════════════
    {
        "id": "UC-23", "name": "View Booking List (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views all bookings with filtering by status.",
        "trigger": "Staff needs to review or process bookings.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: Filtered list of bookings is displayed.",
        "normal_flow": (
            "1. Staff accesses booking list.\n"
            "2. System retrieves bookings.\n"
            "3. Staff optionally filters by status (Pending, Confirmed, Checked-in, etc.).\n"
            "4. System displays matching bookings."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-24", "name": "View Booking Detail (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views complete details of a specific booking.",
        "trigger": "Staff selects a booking to review.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Booking exists.",
        "postconditions": "POST-1: Full booking details displayed (customer info, room, dates, payment, occupants).",
        "normal_flow": (
            "1. Staff selects a booking.\n"
            "2. System retrieves full booking details including customer, room, payment, and occupants.\n"
            "3. System displays all information."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-25", "name": "Assign Room to Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff assigns a specific available room to a confirmed booking.",
        "trigger": "Customer arrives for check-in and booking needs room assignment.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Booking is in Confirmed status.\nPRE-3: Rooms of the booked type are available.",
        "postconditions": "POST-1: Room is assigned to the booking.\nPOST-2: Room status is updated to Occupied.",
        "normal_flow": (
            "1. Staff selects the confirmed booking.\n"
            "2. System displays available rooms of the booked type.\n"
            "3. Staff selects a room to assign.\n"
            "4. System assigns the room and updates room status to Occupied.\n"
            "5. System updates booking status to Checked-in.\n"
            "6. System confirms assignment."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): No rooms of required type available. System notifies staff.",
        "priority": "High", "frequency": "High",
        "business_rules": "BR-1: Only rooms matching the booked room type can be assigned.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-26", "name": "Check Out Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff processes guest checkout, finalizes charges, and frees the room.",
        "trigger": "Guest is leaving and needs to check out.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Booking is in Checked-in status.",
        "postconditions": "POST-1: Booking status updated to Checked-out.\nPOST-2: Room status updated to Cleaning.\nPOST-3: Final invoice is generated.",
        "normal_flow": (
            "1. Staff selects the active booking.\n"
            "2. System calculates final charges (room + extensions + services).\n"
            "3. System generates the final invoice.\n"
            "4. Staff confirms checkout.\n"
            "5. System updates booking status to Checked-out.\n"
            "6. System sets room status to Cleaning."
        ),
        "alternative_flows": "AF-1 (at step 2): Outstanding balance exists. Staff processes remaining payment before checkout.",
        "exceptions": "",
        "priority": "High", "frequency": "High",
        "business_rules": "BR-1: All charges must be settled before checkout.\nBR-2: Room enters Cleaning status after checkout.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-27", "name": "Create Walk-in Booking",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff creates a booking for a walk-in guest who arrives without prior reservation (3-step process).",
        "trigger": "Guest arrives at hotel without a reservation.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Rooms are available.",
        "postconditions": "POST-1: Walk-in booking is created.\nPOST-2: Room is assigned and status set to Occupied.\nPOST-3: Guest is checked in immediately.",
        "normal_flow": (
            "1. Staff initiates walk-in booking process.\n"
            "2. Staff enters guest information (name, phone, ID) or selects existing customer.\n"
            "3. System creates or retrieves customer record.\n"
            "4. Staff selects room type and specific room.\n"
            "5. Staff specifies check-in (now) and expected check-out date.\n"
            "6. System calculates price.\n"
            "7. Staff adds occupant details.\n"
            "8. Staff confirms the booking.\n"
            "9. System creates booking with Checked-in status and assigns room."
        ),
        "alternative_flows": "AF-1 (at step 2): Guest is an existing customer. System auto-fills their information.",
        "exceptions": "EF-1 (at step 4): No rooms available. System notifies staff.",
        "priority": "High", "frequency": "Medium",
        "business_rules": "BR-1: Walk-in bookings are immediately checked-in.\nBR-2: Occupant information is required.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-28", "name": "Record Occupants",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff records occupant details (guests staying in the room) for a booking.",
        "trigger": "Guests are checking in and occupant details need to be recorded.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Booking exists.",
        "postconditions": "POST-1: Occupant details are stored for the booking.",
        "normal_flow": (
            "1. Staff accesses the booking's occupant section.\n"
            "2. Staff enters occupant details (name, ID number, phone).\n"
            "3. System validates the information.\n"
            "4. System saves occupant records.\n"
            "5. Staff can add additional occupants if needed."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 3): Number of occupants exceeds room capacity. System warns.",
        "priority": "Medium", "frequency": "High",
        "business_rules": "BR-1: Number of occupants must not exceed room capacity.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - ROOM MANAGEMENT ═══════════════════
    {
        "id": "UC-29", "name": "View Room Map (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views a visual map showing all rooms and their current status.",
        "trigger": "Staff needs an overview of room availability and statuses.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: Room map with color-coded statuses is displayed.",
        "normal_flow": (
            "1. Staff accesses room map.\n"
            "2. System retrieves all rooms with current status.\n"
            "3. System displays visual map with status indicators (Available, Occupied, Cleaning, Maintenance)."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-30", "name": "View Room Detail (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views detailed information about a specific room including current booking.",
        "trigger": "Staff selects a room from the map or list.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: Room details and current occupancy info displayed.",
        "normal_flow": (
            "1. Staff selects a room.\n"
            "2. System retrieves room details (number, type, status, current booking if occupied).\n"
            "3. System displays the information."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-31", "name": "View Room History (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views the booking history of a specific room.",
        "trigger": "Staff needs to review past usage of a room.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: Room's booking history is displayed.",
        "normal_flow": (
            "1. Staff selects a room and requests history.\n"
            "2. System retrieves all past bookings for the room.\n"
            "3. System displays booking history with dates, guests, and statuses."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - CLEANING ═══════════════════
    {
        "id": "UC-32", "name": "View Cleaning Tasks",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views list of rooms that need cleaning.",
        "trigger": "Staff needs to see which rooms require cleaning.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: List of rooms with cleaning status is displayed.",
        "normal_flow": (
            "1. Staff accesses cleaning task list.\n"
            "2. System retrieves rooms with Cleaning status.\n"
            "3. System displays room numbers and checkout timestamps."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-33", "name": "Update Cleaning Status",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff updates the cleaning status of a room after cleaning is completed.",
        "trigger": "Room cleaning has been completed.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Room is in Cleaning status.",
        "postconditions": "POST-1: Room status is updated to Available.\nPOST-2: Room is ready for new guests.",
        "normal_flow": (
            "1. Staff selects the cleaned room.\n"
            "2. Staff marks cleaning as complete.\n"
            "3. System updates room status to Available.\n"
            "4. System confirms status update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "BR-1: Room transitions from Cleaning to Available status.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - PAYMENT ═══════════════════
    {
        "id": "UC-34", "name": "Process Staff Payment",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff processes payment at the hotel (cash or card) for a booking.",
        "trigger": "Guest needs to pay remaining balance at the hotel.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Booking has outstanding balance.",
        "postconditions": "POST-1: Payment is recorded.\nPOST-2: Invoice is updated.",
        "normal_flow": (
            "1. Staff selects the booking requiring payment.\n"
            "2. System displays outstanding amount.\n"
            "3. Staff selects payment method (cash or VNPay).\n"
            "4. System processes the payment.\n"
            "5. System records the transaction and updates invoice.\n"
            "6. System confirms payment."
        ),
        "alternative_flows": "AF-1 (at step 3): Staff selects VNPay. System generates payment QR/URL for guest.",
        "exceptions": "",
        "priority": "High", "frequency": "High",
        "business_rules": "BR-1: Cash payments are recorded immediately.\nBR-2: VNPay payments require gateway confirmation.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - SERVICE REQUESTS ═══════════════════
    {
        "id": "UC-35", "name": "View Service Requests (Staff)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views pending and in-progress service requests.",
        "trigger": "Staff needs to handle service requests from guests.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: List of service requests is displayed.",
        "normal_flow": (
            "1. Staff accesses service request list.\n"
            "2. System retrieves all service requests.\n"
            "3. System displays requests with type, priority, status, room, and timestamps."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-36", "name": "Update Service Request Status",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff updates the status of a service request (accept, complete, reject).",
        "trigger": "Staff is working on or has completed a service request.",
        "preconditions": "PRE-1: Staff is logged in.\nPRE-2: Service request exists.",
        "postconditions": "POST-1: Service request status is updated.\nPOST-2: Completion time is recorded if completed.",
        "normal_flow": (
            "1. Staff selects a service request.\n"
            "2. Staff updates status (In Progress / Completed) and adds notes.\n"
            "3. System records status change and timestamps.\n"
            "4. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "BR-1: Completed requests record the completion timestamp.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ STAFF - LOGIN ═══════════════════
    {
        "id": "UC-37", "name": "Staff Login",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff authenticates to access the staff portal.",
        "trigger": "Staff member needs to access the staff dashboard.",
        "preconditions": "PRE-1: Staff account exists and is active.",
        "postconditions": "POST-1: Staff session is established.\nPOST-2: Staff is redirected to staff dashboard.",
        "normal_flow": (
            "1. Staff provides email and password.\n"
            "2. System validates credentials and verifies staff role.\n"
            "3. System creates staff session.\n"
            "4. System redirects to staff dashboard."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Invalid credentials. System notifies.\nEF-2 (at step 2): Account not a staff role. System denies access.",
        "priority": "High", "frequency": "Very High",
        "business_rules": "BR-1: Only accounts with Staff role can access staff portal.",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - ROOM MANAGEMENT ═══════════════════
    {
        "id": "UC-38", "name": "Add Room",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin adds a new room to the hotel inventory.",
        "trigger": "Hotel adds a new room.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Room type exists.",
        "postconditions": "POST-1: New room is added with Available status.",
        "normal_flow": (
            "1. Admin provides room number and selects room type.\n"
            "2. System validates room number uniqueness.\n"
            "3. System creates the room with Available status.\n"
            "4. System confirms creation."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Room number already exists. System rejects.",
        "priority": "High", "frequency": "Low",
        "business_rules": "BR-1: Room numbers must be unique.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-39", "name": "Edit Room",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin edits room information (room number, type, status).",
        "trigger": "Room details need to be updated.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Room exists.",
        "postconditions": "POST-1: Room information is updated.",
        "normal_flow": (
            "1. Admin selects a room to edit.\n"
            "2. System displays current room details.\n"
            "3. Admin modifies desired fields.\n"
            "4. System validates and saves changes.\n"
            "5. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 4): Validation fails. System highlights errors.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-40", "name": "View Room List (Admin)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all rooms with status and type information.",
        "trigger": "Admin needs room inventory overview.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Complete room list is displayed.",
        "normal_flow": (
            "1. Admin accesses room list.\n"
            "2. System retrieves all rooms with type and status.\n"
            "3. System displays the list with sorting/filtering options."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-41", "name": "View Room Map (Admin)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views visual room map with status indicators.",
        "trigger": "Admin needs visual overview of room statuses.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Visual room map displayed.",
        "normal_flow": (
            "1. Admin accesses room map.\n"
            "2. System retrieves all rooms and statuses.\n"
            "3. System displays color-coded visual room map."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-42", "name": "View Room History (Admin)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views booking history for a specific room.",
        "trigger": "Admin needs room usage history.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Room booking history displayed.",
        "normal_flow": (
            "1. Admin selects a room.\n"
            "2. System retrieves all past bookings for the room.\n"
            "3. System displays history with dates, guests, statuses."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - ROOM TYPE ═══════════════════
    {
        "id": "UC-43", "name": "Add Room Type",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin creates a new room type with pricing and amenities.",
        "trigger": "Hotel introduces a new category of rooms.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: New room type is created and available for room assignment.",
        "normal_flow": (
            "1. Admin provides type name, base price, hourly price, capacity, deposit percentage, and description.\n"
            "2. Admin selects amenities for this room type.\n"
            "3. Admin uploads images.\n"
            "4. System validates and creates the room type.\n"
            "5. System confirms creation."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 4): Type name already exists. System rejects.",
        "priority": "High", "frequency": "Low",
        "business_rules": "BR-1: Room type names must be unique.\nBR-2: Base price and capacity are required.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-44", "name": "Edit Room Type",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin updates room type details including pricing and amenities.",
        "trigger": "Room type pricing or features need updating.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Room type exists.",
        "postconditions": "POST-1: Room type information is updated.",
        "normal_flow": (
            "1. Admin selects room type to edit.\n"
            "2. System displays current details.\n"
            "3. Admin modifies fields (name, pricing, capacity, amenities, images).\n"
            "4. System validates and saves changes.\n"
            "5. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-45", "name": "View Room Type List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all configured room types.",
        "trigger": "Admin needs to review room type configurations.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: List of room types displayed with pricing and capacity.",
        "normal_flow": (
            "1. Admin accesses room type list.\n"
            "2. System retrieves all room types.\n"
            "3. System displays types with name, price, capacity, and room count."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - STAFF MANAGEMENT ═══════════════════
    {
        "id": "UC-46", "name": "Add Staff",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin creates a new staff account.",
        "trigger": "New staff member joins the hotel.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Staff account is created with Staff role.",
        "normal_flow": (
            "1. Admin provides staff details (name, email, phone, address, password).\n"
            "2. System validates information and email uniqueness.\n"
            "3. System creates account with Staff role.\n"
            "4. System confirms creation."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Email already in use. System rejects.",
        "priority": "High", "frequency": "Low",
        "business_rules": "BR-1: Staff accounts have the Staff role assigned.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-47", "name": "Edit Staff",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin updates staff information or deactivates staff account.",
        "trigger": "Staff information changes or staff leaves.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Staff account exists.",
        "postconditions": "POST-1: Staff account is updated.",
        "normal_flow": (
            "1. Admin selects staff member.\n"
            "2. System displays current staff details.\n"
            "3. Admin modifies details or toggles active status.\n"
            "4. System saves changes.\n"
            "5. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-48", "name": "View Staff List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all staff members.",
        "trigger": "Admin needs to review staff roster.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Staff list is displayed.",
        "normal_flow": (
            "1. Admin accesses staff list.\n"
            "2. System retrieves all staff accounts.\n"
            "3. System displays staff with name, email, phone, and active status."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - CUSTOMER MANAGEMENT ═══════════════════
    {
        "id": "UC-49", "name": "View Customer List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all registered customers.",
        "trigger": "Admin needs to review customer base.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Customer list displayed with account and membership info.",
        "normal_flow": (
            "1. Admin accesses customer list.\n"
            "2. System retrieves all customers.\n"
            "3. System displays customers with name, email, membership level, loyalty points."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-50", "name": "Edit Customer",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin edits customer information or deactivates customer account.",
        "trigger": "Customer information needs updating by admin.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Customer account exists.",
        "postconditions": "POST-1: Customer account is updated.",
        "normal_flow": (
            "1. Admin selects customer.\n"
            "2. System displays customer details.\n"
            "3. Admin modifies information or toggles active status.\n"
            "4. System saves changes.\n"
            "5. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - PROMOTIONS ═══════════════════
    {
        "id": "UC-51", "name": "Add Promotion",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin creates a new promotional campaign.",
        "trigger": "Hotel launches a new promotion.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Promotion is created and active.",
        "normal_flow": (
            "1. Admin provides promotion details (name, description, discount, start/end dates).\n"
            "2. System validates the information.\n"
            "3. System creates the promotion.\n"
            "4. System confirms creation."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): End date before start date. System rejects.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Promotion dates must be valid range.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-52", "name": "Edit Promotion",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin updates or deactivates an existing promotion.",
        "trigger": "Promotion details need changing.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Promotion exists.",
        "postconditions": "POST-1: Promotion is updated.",
        "normal_flow": (
            "1. Admin selects promotion.\n"
            "2. System displays current details.\n"
            "3. Admin modifies fields.\n"
            "4. System validates and saves.\n"
            "5. System confirms."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-53", "name": "View Promotion List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all promotions.",
        "trigger": "Admin reviews active and past promotions.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Promotion list displayed.",
        "normal_flow": (
            "1. Admin accesses promotion list.\n"
            "2. System retrieves all promotions.\n"
            "3. System displays promotions with status, discount, and dates."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - VOUCHERS ═══════════════════
    {
        "id": "UC-54", "name": "Add Voucher",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin creates a new discount voucher for customers.",
        "trigger": "Hotel wants to issue discount vouchers.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Voucher is created and available for use.",
        "normal_flow": (
            "1. Admin provides voucher details (code, discount amount/percentage, expiry, usage limit).\n"
            "2. System validates voucher code uniqueness.\n"
            "3. System creates the voucher.\n"
            "4. System confirms creation."
        ),
        "alternative_flows": "",
        "exceptions": "EF-1 (at step 2): Voucher code already exists. System rejects.",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "BR-1: Voucher codes must be unique.\nBR-2: Voucher must have expiry date.",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-55", "name": "Edit Voucher",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin edits voucher details or deactivates a voucher.",
        "trigger": "Voucher needs updating.",
        "preconditions": "PRE-1: Admin is logged in.\nPRE-2: Voucher exists.",
        "postconditions": "POST-1: Voucher is updated.",
        "normal_flow": (
            "1. Admin selects voucher.\n"
            "2. System displays current details.\n"
            "3. Admin modifies fields.\n"
            "4. System validates and saves.\n"
            "5. System confirms."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-56", "name": "View Voucher List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all vouchers with usage statistics.",
        "trigger": "Admin reviews voucher inventory.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Voucher list displayed.",
        "normal_flow": (
            "1. Admin accesses voucher list.\n"
            "2. System retrieves all vouchers.\n"
            "3. System displays vouchers with code, discount, expiry, and usage count."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - FEEDBACK ═══════════════════
    {
        "id": "UC-57", "name": "View Feedback List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all customer feedback/reviews.",
        "trigger": "Admin reviews customer satisfaction.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Feedback list displayed with ratings and comments.",
        "normal_flow": (
            "1. Admin accesses feedback list.\n"
            "2. System retrieves all feedback.\n"
            "3. System displays feedback with customer name, rating, comments, and date."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - REPORTS ═══════════════════
    {
        "id": "UC-58", "name": "View Revenue Report",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views revenue reports with filtering by date range.",
        "trigger": "Admin needs financial overview.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Revenue report is displayed with charts and totals.",
        "normal_flow": (
            "1. Admin accesses revenue report.\n"
            "2. Admin selects date range.\n"
            "3. System calculates revenue from bookings and payments.\n"
            "4. System displays revenue data with breakdown and totals."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-59", "name": "View Room Utilization Report",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views room utilization/occupancy statistics.",
        "trigger": "Admin needs to analyze room usage efficiency.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Utilization report with occupancy rates displayed.",
        "normal_flow": (
            "1. Admin accesses utilization report.\n"
            "2. Admin selects date range.\n"
            "3. System calculates occupancy rates per room type.\n"
            "4. System displays utilization data."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },

    # ═══════════════════ ADMIN - SETTINGS & INFO ═══════════════════
    {
        "id": "UC-60", "name": "Update Hotel Information",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin updates hotel details displayed on the website (name, address, contact, description).",
        "trigger": "Hotel information needs updating.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Hotel information is updated on the website.",
        "normal_flow": (
            "1. Admin accesses hotel information settings.\n"
            "2. System displays current information.\n"
            "3. Admin modifies details (name, address, phone, email, description).\n"
            "4. System saves changes.\n"
            "5. System confirms update."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Low", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-61", "name": "View Admin Dashboard",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views the main dashboard with key metrics and statistics.",
        "trigger": "Admin accesses the admin panel.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Dashboard with key metrics displayed.",
        "normal_flow": (
            "1. Admin accesses the dashboard.\n"
            "2. System calculates key metrics (total bookings, revenue, occupancy rate, pending requests).\n"
            "3. System displays the dashboard with summary widgets."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-62", "name": "View User List",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all user accounts across all roles.",
        "trigger": "Admin needs to review all system users.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: User list displayed with roles and status.",
        "normal_flow": (
            "1. Admin accesses user list.\n"
            "2. System retrieves all accounts.\n"
            "3. System displays users with name, email, role, and active status."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Low",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-63", "name": "View Service Requests (Admin)",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Admin", "secondary_actor": "System",
        "description": "Admin views all service requests across the hotel for monitoring.",
        "trigger": "Admin needs to monitor service quality.",
        "preconditions": "PRE-1: Admin is logged in.",
        "postconditions": "POST-1: Complete list of service requests displayed.",
        "normal_flow": (
            "1. Admin accesses service request list.\n"
            "2. System retrieves all service requests.\n"
            "3. System displays requests with type, status, staff assigned, and timestamps."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "Medium", "frequency": "Medium",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-64", "name": "View Staff Dashboard",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Staff", "secondary_actor": "System",
        "description": "Staff views operational dashboard with today's tasks and status overview.",
        "trigger": "Staff logs in to begin work.",
        "preconditions": "PRE-1: Staff is logged in.",
        "postconditions": "POST-1: Dashboard with operational summary displayed.",
        "normal_flow": (
            "1. Staff accesses dashboard.\n"
            "2. System calculates today's metrics (check-ins, check-outs, pending requests, rooms to clean).\n"
            "3. System displays dashboard with summary."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
    {
        "id": "UC-65", "name": "View Home Page",
        "created_by": "Team", "date": "25/03/2026",
        "primary_actor": "Guest/Customer", "secondary_actor": "System",
        "description": "User views the hotel website home page with general information and featured rooms.",
        "trigger": "User accesses the hotel website.",
        "preconditions": "",
        "postconditions": "POST-1: Home page with hotel info, featured rooms, and promotions displayed.",
        "normal_flow": (
            "1. User accesses the website.\n"
            "2. System retrieves hotel info, featured room types, and active promotions.\n"
            "3. System displays the home page."
        ),
        "alternative_flows": "",
        "exceptions": "",
        "priority": "High", "frequency": "Very High",
        "business_rules": "",
        "other_info": "", "assumptions": ""
    },
]


# ── Word Document Generator ─────────────────────────────────────────

def set_cell_text(cell, text, bold=False, size=10):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    # Set Vietnamese-compatible font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color
    })
    shading.append(shading_elem)


def merge_cells(table, row, start_col, end_col):
    """Merge cells in a row from start_col to end_col."""
    cell_start = table.cell(row, start_col)
    cell_end = table.cell(row, end_col)
    cell_start.merge(cell_end)


def add_uc_table(doc, uc):
    """Add one UC description table to the document."""
    table = doc.add_table(rows=15, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(3.5)
        row.cells[3].width = Cm(4.5)

    # Row 0: ID and Name (merge cols 1-3)
    merge_cells(table, 0, 1, 3)
    set_cell_text(table.cell(0, 0), "ID and Name", bold=True, size=10)
    set_cell_shading(table.cell(0, 0), "D9E2F3")
    set_cell_text(table.cell(0, 1), f"{uc['id']}: {uc['name']}", bold=True, size=10)

    # Row 1: Created By / Date Created (4 cells)
    set_cell_text(table.cell(1, 0), "Created By", bold=True, size=10)
    set_cell_shading(table.cell(1, 0), "D9E2F3")
    set_cell_text(table.cell(1, 1), uc["created_by"], size=10)
    set_cell_text(table.cell(1, 2), "Date Created", bold=True, size=10)
    set_cell_shading(table.cell(1, 2), "D9E2F3")
    set_cell_text(table.cell(1, 3), uc["date"], size=10)

    # Row 2: Primary Actor / Secondary Actor (4 cells)
    set_cell_text(table.cell(2, 0), "Primary Actor", bold=True, size=10)
    set_cell_shading(table.cell(2, 0), "D9E2F3")
    set_cell_text(table.cell(2, 1), uc["primary_actor"], size=10)
    set_cell_text(table.cell(2, 2), "Secondary Actor", bold=True, size=10)
    set_cell_shading(table.cell(2, 2), "D9E2F3")
    set_cell_text(table.cell(2, 3), uc["secondary_actor"], size=10)

    # Rows 3-14: label + merged value (cols 1-3)
    rows_data = [
        (3, "Description", uc["description"]),
        (4, "Trigger", uc["trigger"]),
        (5, "Preconditions", uc["preconditions"]),
        (6, "Postconditions", uc["postconditions"]),
        (7, "Normal Flow", uc["normal_flow"]),
        (8, "Alternative Flows", uc["alternative_flows"]),
        (9, "Exceptions", uc["exceptions"]),
        (10, "Priority", uc["priority"]),
        (11, "Frequency of Use", uc["frequency"]),
        (12, "Business Rules", uc["business_rules"]),
        (13, "Other Information", uc["other_info"]),
        (14, "Assumptions", uc["assumptions"]),
    ]

    for row_idx, label, value in rows_data:
        merge_cells(table, row_idx, 1, 3)
        set_cell_text(table.cell(row_idx, 0), label, bold=True, size=10)
        set_cell_shading(table.cell(row_idx, 0), "D9E2F3")
        set_cell_text(table.cell(row_idx, 1), value if value else "N/A", size=10)

    # Add spacing after table
    doc.add_paragraph("")


def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Use Case Descriptions - Hotel Management System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    doc.add_paragraph(
        f"Total Use Cases: {len(UC_DATA)}\n"
        "Generated from codebase analysis.\n"
        "Date: 25/03/2026"
    )

    # Group by actor sections
    sections = {
        "Authentication (Guest/Customer)": ["UC-01", "UC-02", "UC-03", "UC-04", "UC-05", "UC-06", "UC-07"],
        "Room Browsing (Guest/Customer)": ["UC-08", "UC-09"],
        "Booking (Customer)": ["UC-10", "UC-11", "UC-12", "UC-13", "UC-14", "UC-15"],
        "Payment (Customer)": ["UC-16"],
        "Customer Profile and Feedback": ["UC-17", "UC-18", "UC-19", "UC-20"],
        "Service Requests (Customer)": ["UC-21", "UC-22"],
        "Staff - Booking Operations": ["UC-23", "UC-24", "UC-25", "UC-26", "UC-27", "UC-28"],
        "Staff - Room Operations": ["UC-29", "UC-30", "UC-31"],
        "Staff - Cleaning": ["UC-32", "UC-33"],
        "Staff - Payment": ["UC-34"],
        "Staff - Service Requests": ["UC-35", "UC-36"],
        "Staff - Authentication": ["UC-37"],
        "Staff - Dashboard": ["UC-64"],
        "Admin - Room Management": ["UC-38", "UC-39", "UC-40", "UC-41", "UC-42"],
        "Admin - Room Type Management": ["UC-43", "UC-44", "UC-45"],
        "Admin - Staff Management": ["UC-46", "UC-47", "UC-48"],
        "Admin - Customer Management": ["UC-49", "UC-50"],
        "Admin - Promotions": ["UC-51", "UC-52", "UC-53"],
        "Admin - Vouchers": ["UC-54", "UC-55", "UC-56"],
        "Admin - Feedback": ["UC-57"],
        "Admin - Reports": ["UC-58", "UC-59"],
        "Admin - Settings and Info": ["UC-60", "UC-61", "UC-62", "UC-63"],
        "Home Page": ["UC-65"],
    }

    uc_map = {uc["id"]: uc for uc in UC_DATA}

    for section_name, uc_ids in sections.items():
        doc.add_heading(section_name, level=1)
        for uc_id in uc_ids:
            if uc_id in uc_map:
                add_uc_table(doc, uc_map[uc_id])

    # Save
    output_path = sys.argv[1] if len(sys.argv) > 1 else "UC-Descriptions-HotelManagementSystem.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")
    print(f"Total UCs: {len(UC_DATA)}")


if __name__ == "__main__":
    main()
